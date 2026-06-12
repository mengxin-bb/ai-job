"""
简历文本提取：从上传的 PDF / DOCX / TXT 文件里抽出纯文本，供 AI 分析。

对外只暴露 extract_text(filename, file_bytes) -> str。
提取失败时抛 ResumeParseError，调用方捕获后给用户友好提示。
"""

from __future__ import annotations

import io


class ResumeParseError(Exception):
    """简历解析失败（格式不支持、文件损坏、内容为空等）。"""


# 单文件大小上限，防止超大文件拖垮内存/接口（5 MB 对简历足够）。
MAX_BYTES = 5 * 1024 * 1024

# 提取后文本上限，避免超长内容撑爆 token（约几千字的简历足够）。
MAX_CHARS = 12000


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    parts = [(page.extract_text() or "") for page in reader.pages]
    return "\n".join(parts)


def _extract_docx(data: bytes) -> str:
    import docx

    doc = docx.Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs]
    # 也提取表格里的文字（很多简历用表格排版）。
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def _extract_txt(data: bytes) -> str:
    for enc in ("utf-8", "gbk", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def extract_text(filename: str, file_bytes: bytes) -> str:
    """
    根据扩展名从简历文件提取纯文本。

    参数:
        filename:   原始文件名（用于判断扩展名）
        file_bytes: 文件二进制内容

    返回:
        提取出的纯文本（已去除多余空白、截断到 MAX_CHARS）。

    异常:
        ResumeParseError: 格式不支持 / 文件过大 / 解析失败 / 内容为空。
    """
    if not file_bytes:
        raise ResumeParseError("文件是空的，请重新上传。")
    if len(file_bytes) > MAX_BYTES:
        raise ResumeParseError("文件太大了（超过 5MB），请压缩或导出更小的版本。")

    name = (filename or "").lower()
    try:
        if name.endswith(".pdf"):
            text = _extract_pdf(file_bytes)
        elif name.endswith(".docx"):
            text = _extract_docx(file_bytes)
        elif name.endswith(".txt"):
            text = _extract_txt(file_bytes)
        elif name.endswith(".doc"):
            raise ResumeParseError(
                "暂不支持旧版 .doc 格式，请用 Word 另存为 .docx，或导出 PDF 后再上传。"
            )
        else:
            raise ResumeParseError("仅支持 PDF / DOCX / TXT 格式的简历。")
    except ResumeParseError:
        raise
    except Exception as e:  # noqa: BLE001 —— 解析库各种异常统一兜底
        raise ResumeParseError(f"文件解析失败：{e}") from e

    # 归一化空白：去掉每行首尾空格、压掉空行。
    lines = [ln.strip() for ln in text.splitlines()]
    cleaned = "\n".join(ln for ln in lines if ln)

    if len(cleaned) < 30:
        raise ResumeParseError(
            "没能从文件里读到足够文字。如果是扫描版/图片 PDF，请上传带文字的版本。"
        )

    return cleaned[:MAX_CHARS]
